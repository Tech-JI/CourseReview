"""Syllabus text extraction and local Ollama analysis helpers.

Pure functions (settings-only dependencies) so they can be unit-tested and
reused from the Celery task in apps/web/tasks.py.
"""

from __future__ import annotations

import base64
import io
import json
import logging
import re

import httpx
from django.conf import settings

logger = logging.getLogger(__name__)

# Below this many characters a PDF is treated as a scan and OCR'd via vision.
TEXT_OCR_THRESHOLD = 300
# Prompts are capped so a long syllabus never blows the context window.
MAX_SYLLABUS_CHARS = 60_000
MAX_SUMMARY_CHARS = 2_000


class SyllabusAnalysisError(Exception):
    """Raised when extraction, OCR or model analysis fails."""


def extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract embedded text from a PDF; '' for scanned/image-only PDFs."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "\n".join(
        (page.extract_text() or "") for page in reader.pages if page.extract_text()
    )


def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract paragraphs + table cells from a .docx file."""
    import docx

    document = docx.Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def render_pdf_pages(pdf_bytes: bytes, max_pages: int | None = None) -> list[bytes]:
    """Render PDF pages to PNG bytes for vision OCR (no system deps)."""
    import pypdfium2 as pdfium

    max_pages = max_pages or settings.OLLAMA["MAX_PAGES"]
    pdf = pdfium.PdfDocument(pdf_bytes)
    try:
        n_pages = min(len(pdf), max_pages)
        images: list[bytes] = []
        for i in range(n_pages):
            page = pdf[i]
            bitmap = page.render(scale=150 / 72)  # ~150 dpi
            pil_image = bitmap.to_pil()
            buffer = io.BytesIO()
            pil_image.save(buffer, format="PNG")
            images.append(buffer.getvalue())
            pil_image.close()
        return images
    finally:
        pdf.close()


def _image_message(page_png: bytes, page_no: int) -> dict:
    encoded = base64.b64encode(page_png).decode("ascii")
    return {
        "type": "text",
        "text": f"Page {page_no}:",
    }, {
        "type": "image_url",
        "image_url": {"url": f"data:image/png;base64,{encoded}"},
    }


def ocr_pages(images: list[bytes]) -> str:
    """OCR rendered pages with the local vision model."""
    if not images:
        return ""
    parts: list[dict] = [
        {"type": "text", "text": "Transcribe all text in each page exactly, in order."}
    ]
    for i, page_png in enumerate(images, start=1):
        parts.extend(_image_message(page_png, i))
    response = ollama_chat(
        [{"role": "user", "content": parts}],
        format_json=False,
    )
    return (response.get("message") or {}).get("content", "").strip()


def ollama_chat(messages: list[dict], format_json: bool = False) -> dict:
    """POST /api/chat to the local Ollama; returns the parsed response dict.

    Retries transient connection/read-timeout failures with backoff.
    """
    ollama = settings.OLLAMA
    payload: dict = {
        "model": ollama["MODEL"],
        "messages": messages,
        "stream": False,
        "options": {
            "num_ctx": ollama["NUM_CTX"],
            "temperature": 0.2,
            # qwen3 thinking models: with JSON format enabled the final answer
            # lands in `message.thinking` and `content` comes back empty.
            "think": False,
        },
    }
    if format_json:
        payload["format"] = "json"
    url = f"{ollama['BASE_URL'].rstrip('/')}/api/chat"
    timeout = ollama["TIMEOUT"]
    last_error: Exception | None = None
    for attempt, backoff in enumerate((0, 5, 15)):
        if backoff:
            import time

            time.sleep(backoff)
        try:
            with httpx.Client(timeout=timeout) as client:
                response = client.post(url, json=payload)
            response.raise_for_status()
            data = response.json()
            if not (data.get("message") or {}).get("content"):
                logger.warning(
                    "Ollama returned empty content: done=%s eval=%d ctx=%d model=%s",
                    data.get("done_reason"),
                    data.get("eval_count"),
                    data.get("prompt_eval_count"),
                    payload["model"],
                )
            return data
        except (httpx.ConnectError, httpx.ReadTimeout) as exc:
            last_error = exc
            logger.warning("Ollama unreachable (attempt %d/3): %s", attempt + 1, exc)
        except httpx.HTTPStatusError as exc:
            raise SyllabusAnalysisError(
                f"Ollama returned HTTP {exc.response.status_code}"
            ) from exc
    raise SyllabusAnalysisError(f"Ollama unreachable: {last_error}")


def parse_json_response(content: str) -> dict:
    """Parse model output to dict: strip fences, strict parse, regex fallback."""
    if not content:
        raise SyllabusAnalysisError("Empty model response")
    text = content.strip()
    text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.MULTILINE)
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        # Not clean JSON yet; fall through to fence-stripping / regex below.
        pass
    match = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            # Regex extraction failed too; the final raise reports the failure.
            pass
    raise SyllabusAnalysisError("Model output was not parseable JSON")


def _truncate(text: str, limit: int) -> str:
    return text if len(text) <= limit else text[:limit] + "\n...[truncated]"


def build_analysis_prompt(course, instructor, syllabus_text: str) -> str:
    """Prompt asking the model to check the syllabus against known course info."""
    topics = course.course_topics or []
    topic_block = (
        "\n".join(f"- {topic}" for topic in topics)
        if topics
        else "(no known topics on file)"
    )
    return f"""You are verifying a course syllabus for authenticity against known course data.

Course code: {course.course_code}
Course title: {course.course_title}
Department: {course.department}
Known description: {course.description or "(none)"}
Known topics on file:
{topic_block}
Instructor (claimed): {instructor.name}

Syllabus text to analyze:
---
{syllabus_text}
---

Task:
1. Judge whether this syllabus plausibly belongs to this course and instructor (match).
2. Judge whether it looks like a real, legitimate syllabus (not fabricated, clearly
   copied from another course, or garbled).
3. Produce a concise markdown summary of the syllabus: structure, grading scheme,
   schedule highlights, textbooks, policies.

Reply with JSON only, shape:
{{"match_score": <int 0-100>, "matches_course_content": <bool>, "is_legitimate": <bool>,
"flags": [<string reasons for low score / low legitimacy>], "summary_md": "<markdown summary>"}}
"""


def build_comparison_prompt(course, instructor, new_text: str, old_text: str) -> str:
    return f"""Two syllabus versions exist for {course.course_code} - {course.course_title}
taught by {instructor.name}. Decide which one is newer, matches the course better,
and is more authentic/complete. Existing (primary) version first, new upload second.

Existing version:
---
{old_text}
---

New version:
---
{new_text}
---

Reply with JSON only, shape:
{{"newer": <bool>, "better_match": <bool>, "more_authentic": <bool>,
"recommendation": "keep_old"|"keep_new", "notes": "<short rationale>"}}
"""


def analyze(course, instructor, syllabus_text: str) -> dict:
    """Run the analysis prompt and return the structured verdict dict."""
    prompt = build_analysis_prompt(
        course, instructor, _truncate(syllabus_text, MAX_SYLLABUS_CHARS)
    )
    response = ollama_chat([{"role": "user", "content": prompt}], format_json=True)
    verdict = parse_json_response((response.get("message") or {}).get("content") or "")
    # summary_md is nested inside the verdict JSON by the model.
    return verdict


def compare(course, instructor, new_text: str, old_text: str) -> dict:
    """Run the comparison prompt; returns the comparison dict."""
    prompt = build_comparison_prompt(
        course,
        instructor,
        _truncate(new_text, MAX_SYLLABUS_CHARS),
        _truncate(old_text, MAX_SYLLABUS_CHARS),
    )
    response = ollama_chat([{"role": "user", "content": prompt}], format_json=True)
    return parse_json_response((response.get("message") or {}).get("content") or "")
